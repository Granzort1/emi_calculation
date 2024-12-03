from docx import Document
import openpyxl


def convert_word_tables_to_excel(word_file_path, excel_file_path):
    # Word 문서 로드
    doc = Document(word_file_path)

    # 새 Excel 워크북 생성
    wb = openpyxl.Workbook()
    ws = wb.active

    # 현재 Excel 행 위치 추적
    current_row = 1

    # 첫 번째 표의 헤더 저장
    if doc.tables:
        header_row = doc.tables[0].rows[0]
        headers = [cell.text.strip() for cell in header_row.cells]
        header_dict = {header: idx for idx, header in enumerate(headers)}

        # 첫 번째 표 처리
        first_table = doc.tables[0]
        # 헤더 쓰기
        for col_idx, header in enumerate(headers):
            ws.cell(row=current_row, column=col_idx + 1, value=header)
        current_row += 1

        # 데이터 행 쓰기
        for row in first_table.rows[1:]:
            for col_idx, cell in enumerate(row.cells):
                ws.cell(row=current_row, column=col_idx + 1, value=cell.text.strip())
            current_row += 1

        # 나머지 표 처리
        for table in doc.tables[1:]:
            # 현재 표의 헤더 가져오기
            current_headers = [cell.text.strip() for cell in table.rows[0].cells]
            # 열 매핑 생성
            column_mapping = {}
            for idx, header in enumerate(current_headers):
                if header in header_dict:
                    column_mapping[idx] = header_dict[header]

            # 데이터 행 처리
            for row in table.rows[1:]:
                # 새로운 행 데이터 초기화
                new_row_data = [''] * len(headers)
                # 각 셀의 데이터를 올바른 열에 매핑
                for col_idx, cell in enumerate(row.cells):
                    if col_idx in column_mapping:
                        new_col_idx = column_mapping[col_idx]
                        new_row_data[new_col_idx] = cell.text.strip()

                # 행 데이터 쓰기
                for col_idx, value in enumerate(new_row_data):
                    ws.cell(row=current_row, column=col_idx + 1, value=value)
                current_row += 1

    # Excel 파일 저장
    wb.save(excel_file_path)
    print(f"변환 완료: {excel_file_path}")


# 사용 예시
word_file_path = r'C:\dense_traffic_emi\input\1.docx'
excel_file_path = r'C:\dense_traffic_emi\output\wordto_excel.xlsx'
convert_word_tables_to_excel(word_file_path, excel_file_path)